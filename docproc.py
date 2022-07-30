import docx #you need to install python-docx
import re
import pickle
import os
import random
from datetime import datetime
#define a object that contain a data entry
class Entry:
    def __init__(self):
        self.syno=[] #synonym
        self.anto=[] #antonym
        self.key="" #the word
        self.trans=[] #meaning
        self.type=[] #verb, non, adj..
    def __init__(self, key):
        self.syno=[] #synonym
        self.anto=[] #antonym
        self.key= key #the word
        self.trans=[] #meaning
        self.type=[] #verb, non, adj...

state_flag=False #mark the start of vocabulary or not
read_flag=False #mark the dictionary file already exist or not
count=0 #number of vocabulary
my_dict=[] #to save every parsed entry
#try to read the parsed vocabulary, if extist, do not need to parse again
if os.path. exists("my_dict.obj"):
    with open('my_dict.obj','rb') as dict_file:
        my_dict = pickle.load(dict_file)
        dict_file.close()
        count=len(my_dict)
        read_flag = True
        print("dictionary file already exist, Loading done!")
#if dictinoary is not exist, procesing
if read_flag is not True:
    print("dictionary file does not exist, parsing ...")
    doc = docx.Document('单词书成品可编辑版.docx')
    reg_list=[
        r"^(?!【释义)(.*)\s*\[", #match each vocabulary key, 
        "释义\d?\s*】(.*\.)(.*)", #match one trans with one type (eg vt.)
        "释义\d?\s*】(.*)", #match one trans without type (eg vt.)
        "同义\s*】(.*)",
        "反义\s*】(.*)",
    ]
    e=None
    #go through word document line by line
    for i in doc.paragraphs:
        if re.match(r"List \d*", i.text):
            state_flag = True #word start
            print ("Processing "+ i.text)
            if e is not None:
                my_dict[-1].append(e)
                count+=1
                e = None
            my_dict.append([])
            continue
        if state_flag:
            #match vocabulary key
            m = re.match(reg_list[0], i.text)
            if m:
                #save the previous entry
                if e is not None:
                    my_dict[-1].append(e)
                    count+=1
                #start new entry
                key = m.group(1).strip()
                e = Entry(key)
                continue
            #match each item in reg_list
            for j in range(1, len(reg_list)) :
                m = re.findall(reg_list[j],i.text)
                if m:
                    if j==1:
                        e.trans.append(m[0][1].strip())
                        e.type.append(m[0][0].strip())
                    elif j==2:
                        e.trans.append(m[0].strip())
                        e.type.append ("n/a")
                    elif j==3:
                        e.syno.append(m[0].strip())
                    else:
                        e.anto.append (m[0]. strip())
                    break
    #the last one is not saved yet, so just add the last one
    if e is not None:
        my_dict[-1].append(e)
        count+=1
        with open('my_dict.obj', 'wb') as dict_file:
            pickle.dump (my_dict, dict_file)
            dict_file.close()


random.seed(datetime.now().microsecond)
#generate questions
num=30 #number of questions
questions=[]
answers=[]

for n in range(0, len(my_dict)):
    count = len(my_dict[n])
    sample = random.sample(range(0, count), num) #random select 30 questions
    questions.append([])
    answers.append([])
    # i is the correct key
    #here, 3 error keys are picked
    # a = 1 - 1 is selected, if 1-1 is negative number, just use the last entry in the dictionary
    # b = 1 + 1 is selected, if 1+1 exceeds last entry, just use the first entry in the dictionary
    #j is another random key which is not equal to a, b and i
    # Note, in some cases, more than one translation available for the entry, just random pick one
    for i in sample:
        key = my_dict[n][i] #correct key
        a= (i+random.randint(1,5))%count #correct key +1
        b= (i-random.randint(1,5))%count #correct key -1
        l= len(key.trans)
        c = key.trans[random.randint(0, l-1)] #if more than one translate, random pick one
        while True:
            j= random.randint(0, count-1) #random key
            if j!=i and j!=a and j!=b:
                break
        a = my_dict[n][a].trans
        b = my_dict[n][b].trans
        j = my_dict[n][j].trans
        a = a[random.randint(0, len(a)-1)] #if more than one translate, random pick one
        b = b[random.randint(0, len(b)-1)] #if more than one translate, random pick one
        j = j[random.randint(0, len(j)-1)] #if more than one translate, random pick one
        q = [c,b, a, j]
        random.shuffle(q) #shuffle the choices
        q.append(key.key) #add the key (vocabulary entry) to the last one
        questions[-1].append(q)
        answers[-1].append(q.index(c))



print ("Generating questions ... done!")



#now output the questions into docx
for n in range(0, len(questions)):
    doc_out = docx.Document()
    #doc_out.add_heading('GRE Vocabulary Test', 0) #this is title
    choice=["A.","B.","C.", "D."]
    #p = doc_out.add_paragraph('There are 4 choices in each question, please select the best ONE')
    #doc_out.add_page_break()
    for i in range(0, num):
        q = questions[n][i]
        P= doc_out.add_paragraph('Question '+str(i+1)+': '+ q[-1].upper()) #-1 means last one
        for j in range(0, 4) :
            p = doc_out.add_paragraph()
            if j== answers[n][i]:
                p.add_run('\t'+choice[j]+' '+q[j]).bold = True # make answer  bold font
            else:
                p.add_run('\t'+choice[j]+' '+q[j])
        if (i+1)%4==0:
            doc_out.add_page_break()

    out_fname = 'Question'+str(n)+'.docx'
    doc_out.save(out_fname)
print ("questions saved")


